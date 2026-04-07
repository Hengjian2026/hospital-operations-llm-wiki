#!/usr/bin/env python3
"""
Example ingestion script for the Hospital Operations Knowledge Base.
Demonstrates how to read a .docx file, extract text, and use the Anthropic API
to generate a Markdown summary for the wiki.
"""

import os
import sys
import argparse
import datetime
from pathlib import Path

try:
    import docx
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

try:
    from anthropic import Anthropic
except ImportError:
    print("Please install anthropic: pip install anthropic")
    sys.exit(1)

def extract_text_from_docx(file_path):
    """Extract all text from a .docx file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also try to extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

def generate_wiki_markdown(filename, text, api_key):
    """Call the Anthropic API to generate a wiki Markdown page."""
    client = Anthropic(api_key=api_key)

    prompt = f"""
You are an expert hospital operations data analyst and knowledge base maintainer.
I will provide you with the extracted text from a document titled "{filename}".

Your task is to analyze the content and generate a Markdown summary page for a knowledge base.
Follow the CLAUDE.md schema for the Hospital Operations Knowledge Base.

The output MUST be valid Markdown and include:
1. YAML frontmatter with tags (e.g., #门诊, #DRG, #运营分析).
2. "来源信息" (Source Information): Original filename, date of ingestion, and main topic.
3. "核心内容摘要" (Core Summary): Bullet points of the main findings and conclusions.
4. "关键数据" (Key Metrics): Extract any important metrics, numbers, or tables mentioned.
5. "相关链接" (Related Links): Identify and link to relevant [[概念]] (Concepts), [[实体]] (Entities like departments/diseases), and [[方法论]] (Methodologies). Use Obsidian-style wikilinks.

Document Text:
---
{text[:100000]} # Truncating to avoid token limits in this example
---

Please output ONLY the Markdown content for the wiki page, without any introductory or concluding remarks.
"""

    try:
        response = client.messages.create(
            model="claude-3-opus-20240229",
            max_tokens=4000,
            temperature=0.2,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.content[0].text
    except Exception as e:
        print(f"Error calling Anthropic API: {e}")
        return None

def process_file(file_path, api_key):
    """Process a single docx file and save the result."""
    path = Path(file_path)
    filename = path.name
    print(f"Processing: {filename}")

    text = extract_text_from_docx(file_path)
    if not text:
        return False

    print(f"Extracted {len(text)} characters. Generating summary...")
    markdown_content = generate_wiki_markdown(filename, text, api_key)

    if not markdown_content:
        return False

    # Generate output path
    today = datetime.datetime.now().strftime("%Y%m%d")
    base_name = path.stem
    out_filename = f"{today}-{base_name}.md"

    # Ensure output directory exists based on project structure
    script_dir = Path(__file__).parent.absolute()
    out_dir = script_dir / "wiki" / "来源"
    os.makedirs(out_dir, exist_ok=True)

    out_path = out_dir / out_filename

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(markdown_content)

    print(f"Successfully created: {out_path}")
    return True

def main():
    parser = argparse.ArgumentParser(description="Ingest .docx files into the wiki.")
    parser.add_argument("file", nargs="?", help="Specific .docx file to process")
    parser.add_argument("--dir", default="raw/docx", help="Directory containing .docx files (default: raw/docx)")
    args = parser.parse_args()

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("Warning: ANTHROPIC_API_KEY environment variable is not set.")
        print("The script will fail when trying to call the API.")

    if args.file:
        if not args.file.endswith(".docx"):
            print("Error: File must be a .docx file")
            sys.exit(1)
        process_file(args.file, api_key)
    else:
        # Resolve dir relative to the script location
        script_dir = Path(__file__).parent.absolute()
        dir_path = script_dir / args.dir

        if not dir_path.exists():
            print(f"Error: Directory {dir_path} does not exist.")
            sys.exit(1)

        docx_files = list(dir_path.glob("*.docx"))
        if not docx_files:
            print(f"No .docx files found in {dir_path}")
            return

        print(f"Found {len(docx_files)} .docx files.")
        for f in docx_files:
            process_file(f, api_key)

if __name__ == "__main__":
    main()